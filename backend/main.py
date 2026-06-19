"""
Anderson Report Automation – Self-Hosted Backend
FastAPI server for TERA and PGT-A report generation.

Directory layout expected next to this file:
  server/
    main.py          ← this file
    tera_template.py
    pgta_template.py
    pgta_docx_generator.py
    pgta_classify.py
    tera_assets.py
    pgta_assets.py
    report_comparator.py
    assets/pgta/fonts/   ← fonts used by PDF generators
    reports/             ← TERA output (auto-created)
    reports-pgta/        ← PGTA output (auto-created)
    temp/                ← preview PDFs (auto-created)
    drafts/PGTA/         ← draft JSON files (auto-created)
    uploads/pgta_cnv/    ← uploaded CNV images (auto-created)

Frontend (HTML/JS/CSS) is served from the parent directory:
  ../index.html
  ../tera.html
  ../pgta.html
  ../xlsx.full.min.js
  ../Head_logo.jpg
"""

import pdfplumber
from pgta_classify import auto_map_cnvs
from pgta_docx_generator import PGTADocxGenerator
from pgta_template import PGTAReportTemplate
from tera_template import TERAReportGenerator
import os
import io
import re
import uuid
import json
import math
import base64
import shutil
import difflib
import asyncio
from datetime import datetime
from typing import List

from fastapi import FastAPI, UploadFile, File, Request, BackgroundTasks, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from dotenv import load_dotenv
import pandas as pd

load_dotenv()

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FRONTEND_DIR = os.environ.get("FRONTEND_DIR") or os.path.join(
    os.path.dirname(BASE_DIR), "frontend")

REPORT_DIR = os.path.join(BASE_DIR, "reports")
PGTA_REPORT_DIR = os.path.join(BASE_DIR, "reports-pgta")
TEMP_DIR = os.path.join(BASE_DIR, "temp")
PGTA_CNV_DIR = os.path.join(BASE_DIR, "uploads", "pgta_cnv")
PGTA_DRAFT_DIR = os.path.join(BASE_DIR, "drafts", "PGTA")

for d in (REPORT_DIR, PGTA_REPORT_DIR, TEMP_DIR, PGTA_CNV_DIR, PGTA_DRAFT_DIR):
    os.makedirs(d, exist_ok=True)

# ── Import report generators ───────────────────────────────────────────────────

try:
    from report_comparator import PGTAReportComparator
    _comparator_ok = True
except Exception:
    _comparator_ok = False

# ── NIPT report generators ─────────────────────────────────────────────────────
try:
    from nipt_template import NIPTReportTemplate
    from nipt_docx_generator import NIPTDocxGenerator
    _nipt_ok = True
except Exception as _nipt_err:
    _nipt_ok = False
    print(f"[NIPT] generators not loaded: {_nipt_err}")

# ── HLA Typing report router ───────────────────────────────────────────────────
try:
    from hla_api import router as hla_router
    _hla_ok = True
except Exception as _hla_err:
    _hla_ok = False
    print(f"[HLA] router not loaded: {_hla_err}")

# ── MySQL database ─────────────────────────────────────────────────────────────
_mysql_enabled = False
upload_pdf = upload_pgta_file = save_report = None
try:
    from mysql_client import mysql_enabled as _mysql_enabled, upload_pdf, save_report, upload_pgta_file
except Exception:
    pass

# ── App ────────────────────────────────────────────────────────────────────────
app = FastAPI(title="Anderson Report Automation", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    # covers file:// pages (Origin: null) and all localhost variants
    allow_origin_regex=r".*",
    allow_credentials=False,    # must be False when allow_origins=["*"]
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Static mounts ──────────────────────────────────────────────────────────────
app.mount("/reports",      StaticFiles(directory=REPORT_DIR),      name="reports")
app.mount("/reports-pgta", StaticFiles(directory=PGTA_REPORT_DIR),
          name="reports-pgta")

_fonts_dir = os.path.join(BASE_DIR, "assets", "pgta", "fonts")
_assets_dir = os.path.join(BASE_DIR, "assets", "pgta")
if os.path.isdir(_fonts_dir):
    app.mount("/pgta-fonts",  StaticFiles(directory=_fonts_dir),
              name="pgta-fonts")
if os.path.isdir(_assets_dir):
    app.mount("/pgta-assets", StaticFiles(directory=_assets_dir),
              name="pgta-assets")

# Serve frontend static assets (logo, xlsx.js, css) from frontend dir
if os.path.isdir(FRONTEND_DIR):
    app.mount("/static", StaticFiles(directory=FRONTEND_DIR), name="static")


# ══════════════════════════════════════════════════════════════════════════════
# HTML PAGE ROUTES
# ══════════════════════════════════════════════════════════════════════════════

def _serve(filename: str):
    p = os.path.join(FRONTEND_DIR, filename)
    if os.path.exists(p):
        return FileResponse(p, media_type="text/html")
    raise HTTPException(404, f"{filename} not found")


@app.api_route("/",          methods=["GET", "HEAD"])
@app.api_route("/login.html", methods=["GET", "HEAD"])
@app.api_route("/login",      methods=["GET", "HEAD"])
def login_page(): return _serve("login.html")


@app.get("/otp.html")
@app.get("/otp")
def otp_page(): return _serve("otp.html")


@app.get("/index.html")
@app.get("/home")
def root(): return _serve("index.html")


@app.get("/dashboard")
def tera_page(): return _serve("tera.html")


@app.get("/tera")
@app.get("/tera.html")
def tera_page2(): return _serve("tera.html")


@app.get("/pgta")
@app.get("/pgta.html")
def pgta_page(): return _serve("pgta.html")


@app.get("/nipt")
@app.get("/nipt.html")
def nipt_page(): return _serve("nipt.html")


@app.get("/billing")
@app.get("/billing.html")
def billing_page(): return _serve("billing.html")

# Forward /xlsx.full.min.js  so the frontend can load SheetJS via relative path


@app.get("/xlsx.full.min.js")
def xlsx_js():
    p = os.path.join(FRONTEND_DIR, "xlsx.full.min.js")
    if os.path.exists(p):
        return FileResponse(p, media_type="application/javascript")
    raise HTTPException(404, "xlsx.full.min.js not found")


@app.get("/styles.css")
def styles_css():
    p = os.path.join(FRONTEND_DIR, "styles.css")
    if os.path.exists(p):
        return FileResponse(p, media_type="text/css")
    raise HTTPException(404, "styles.css not found")


@app.get("/Head_logo.jpg")
def head_logo():
    p = os.path.join(FRONTEND_DIR, "Head_logo.jpg")
    if os.path.exists(p):
        return FileResponse(p, media_type="image/jpeg")
    raise HTTPException(404, "Head_logo.jpg not found")


@app.get("/Sign_sec.png")
def sign_sec_png():
    p = os.path.join(FRONTEND_DIR, "Sign_sec.png")
    if os.path.exists(p):
        return FileResponse(p, media_type="image/png")
    raise HTTPException(404, "Sign_sec.png not found")


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _safe_name(name: str) -> str:
    return re.sub(r'[^a-zA-Z0-9 ]', '', str(name).strip()).replace(' ', '_')


def _biopsy_ordinal(biopsy_no: str) -> str:
    m = re.search(r'(\d+)', str(biopsy_no))
    n = int(m.group(1)) if m else 1
    sfx = {1: "st", 2: "nd", 3: "rd"}.get(n if n < 20 else n % 10, "th")
    return f"{n}{sfx} biopsy"


def _build_tera_filename(row: dict, with_logo: bool) -> str:
    patient = _safe_name(row.get("Patient Name", "Unknown"))
    biopsy = _biopsy_ordinal(row.get("Biopsy No.", "1"))
    logo = "with_logo" if with_logo else "without_logo"
    return f"TERA/{patient}_{biopsy}_TERA_report_{logo}.pdf"


def _resolve_cnv_images(embryos: list) -> tuple:
    """Decode any base64 CNV images to temp files so the PDF generator can read them."""
    tmp_paths = []
    for emb in embryos:
        b64 = (emb.get("cnv_image_b64") or "").strip()
        if not b64:
            continue
        try:
            if "," in b64:
                b64 = b64.split(",", 1)[1]
            img_bytes = base64.b64decode(b64)
            tmp_name = f"cnv_{uuid.uuid4().hex}.png"
            tmp_path = os.path.join(TEMP_DIR, tmp_name)
            with open(tmp_path, "wb") as f:
                f.write(img_bytes)
            emb["cnv_image_path"] = tmp_path
            tmp_paths.append(tmp_path)
        except Exception as exc:
            print(f"[cnv_resolve] embryo '{emb.get('embryo_id', '?')}': {exc}")
    return embryos, tmp_paths


def _upload_in_background(filepath: str, filename: str):
    try:
        if upload_pgta_file:
            upload_pgta_file(filepath, filename)
    except Exception as e:
        print(f"[background upload] {filename}: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# TERA ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/preview")
async def preview_report(data: dict):
    file_id = str(uuid.uuid4()) + ".pdf"
    filepath = os.path.join(TEMP_DIR, file_id)
    with_logo = data.get("logo_option", "without_logo") == "with_logo"
    gen = TERAReportGenerator(data, TEMP_DIR, with_logo=with_logo)
    gen.filepath = filepath
    gen.filename = file_id
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, lambda: gen.generate(pages=3))
    return {"preview_url": f"/preview-file/{file_id}"}


@app.get("/preview-file/{filename}")
def preview_file(filename: str):
    path = os.path.join(TEMP_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(404, "Preview not found")
    return FileResponse(path, media_type="application/pdf")


@app.post("/generate")
async def generate_report(data: dict):
    try:
        with_logo = data.get("logo_option", "without_logo") == "with_logo"
        file_name = _build_tera_filename(data, with_logo)
        gen = TERAReportGenerator(data, REPORT_DIR, with_logo=with_logo)
        pdf_path = gen.generate()
        if not pdf_path or not os.path.exists(pdf_path):
            return {"error": "PDF not generated"}
        file_url = upload_pdf(
            pdf_path, file_name) if upload_pdf else f"/reports/{os.path.basename(pdf_path)}"
        return {"status": "success", "file_url": file_url}
    except Exception as e:
        return {"error": str(e)}


@app.post("/generate-bulk")
async def generate_bulk(request: Request):
    data = await request.json()
    output_files, errors = [], []
    for row in data:
        patient_name = row.get("Patient Name", "Unknown")
        try:
            with_logo = row.get("logo_option", "without_logo") == "with_logo"
            file_name = _build_tera_filename(row, with_logo)
            gen = TERAReportGenerator(row, REPORT_DIR, with_logo=with_logo)
            pdf_path = gen.generate()
            file_url = upload_pdf(
                pdf_path, file_name) if upload_pdf else f"/reports/{os.path.basename(pdf_path)}"
            output_files.append(
                {"file_name": os.path.basename(pdf_path), "file_url": file_url})
        except Exception as e:
            errors.append({"patient": patient_name, "error": str(e)})
    return {"generated": output_files, "errors": errors}


@app.post("/upload-excel")
async def upload_excel(file: UploadFile = File(...)):
    try:
        df = pd.read_excel(file.file)

        def _safe(v):
            if v is None:
                return None
            try:
                if pd.isna(v):
                    return None
            except (TypeError, ValueError):
                pass
            if hasattr(v, "item"):
                return v.item()
            if hasattr(v, "isoformat"):
                return str(v)
            return v

        rows = [{k: _safe(v) for k, v in r.items()}
                for r in df.to_dict(orient="records")]
        return {"rows": rows}
    except Exception as e:
        return {"error": str(e), "rows": []}


# ── TERA Draft (file-based, no Supabase required) ─────────────────────────────

TERA_DRAFT_DIR = os.path.join(BASE_DIR, "drafts", "TERA")
os.makedirs(TERA_DRAFT_DIR, exist_ok=True)


@app.post("/save-draft/{draft_type}")
async def save_draft(draft_type: str, request: Request):
    data = await request.json()
    fname = re.sub(r'[^a-zA-Z0-9_-]', '_', draft_type) + ".json"
    with open(os.path.join(TERA_DRAFT_DIR, fname), "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    return {"status": "saved"}


@app.get("/list-drafts")
def list_drafts():
    try:
        files = sorted(os.listdir(TERA_DRAFT_DIR), reverse=True)
        return {"drafts": [{"draft_type": f.replace(".json", ""), "filename": f} for f in files if f.endswith(".json")]}
    except Exception as e:
        return {"drafts": [], "error": str(e)}


@app.get("/load-draft/{draft_type}")
def load_draft(draft_type: str):
    fname = re.sub(r'[^a-zA-Z0-9_-]', '_', draft_type) + ".json"
    fpath = os.path.join(TERA_DRAFT_DIR, fname)
    if not os.path.exists(fpath):
        return {"data": None}
    with open(fpath, encoding="utf-8") as f:
        return {"data": json.load(f)}


# ── TERA Compare ───────────────────────────────────────────────────────────────

@app.post("/tera/compare")
async def tera_compare(manual: UploadFile = File(...), automated: UploadFile = File(...)):
    def extract(data, fname):
        if fname.lower().endswith(".pdf"):
            try:
                lines = []
                with pdfplumber.open(io.BytesIO(data)) as pdf:
                    for pg in pdf.pages:
                        t = pg.extract_text()
                        if t:
                            lines.extend(t.splitlines())
                return [l.strip() for l in lines if l.strip()]
            except Exception as e:
                return [f"[PDF error: {e}]"]
        elif fname.lower().endswith(".docx"):
            try:
                from docx import Document
                doc = Document(io.BytesIO(data))
                return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            except Exception as e:
                return [f"[DOCX error: {e}]"]
        return ["[Unsupported format]"]

    mb, ab = await manual.read(), await automated.read()
    ml, al = extract(mb, manual.filename), extract(ab, automated.filename)
    d = difflib.HtmlDiff(wrapcolumn=80)
    html = d.make_table(ml, al, fromdesc=f"Manual — {manual.filename}",
                        todesc=f"Automated — {automated.filename}", context=True, numlines=3)
    changes = [(t, i1, i2, j1, j2) for t, i1, i2, j1, j2 in difflib.SequenceMatcher(
        None, ml, al).get_opcodes() if t != "equal"]
    return {"html_diff": html, "total_changes": len(changes), "match": len(changes) == 0,
            "manual_file": manual.filename, "auto_file": automated.filename}


# ── Compare PDFs ───────────────────────────────────────────────────────────────


def _norm(s): return re.sub(r'\s+', ' ', s).strip()


def _word_diff(a, b):
    sm = difflib.SequenceMatcher(None, a.split(), b.split(), autojunk=False)
    return [((" ".join(a.split()[i1:i2])), (" ".join(b.split()[j1:j2])))
            for tag, i1, i2, j1, j2 in sm.get_opcodes() if tag != "equal"]


@app.post("/compare-pdf")
async def compare_pdf(file1: UploadFile = File(...), file2: UploadFile = File(...)):
    d1, d2 = await file1.read(), await file2.read()
    sections = []
    with pdfplumber.open(io.BytesIO(d1)) as ld, pdfplumber.open(io.BytesIO(d2)) as rd:
        n1, n2 = len(ld.pages), len(rd.pages)
        pc_msg = f"Manual: {n1} pages — Automated: {n2} pages." if n1 != n2 else f"Both PDFs have {n1} pages. ✓"
        sections.append(("Page Count", [pc_msg]))
        total_pages = max(n1, n2)
        total_diffs = 0
        for i in range(total_pages):
            lt = _norm(ld.pages[i].extract_text() or "") if i < n1 else ""
            rt = _norm(rd.pages[i].extract_text() or "") if i < n2 else ""
            if lt == rt:
                issues = ["Full page text is identical. ✓"]
            else:
                diff_pairs = _word_diff(lt, rt)
                total_diffs += len(diff_pairs)
                issues = [f"<span style='color:#c0392b'>{len(diff_pairs)} difference(s)</span>"] + [
                    f"  <tt>Manual:</tt> <span style='background:#fde8e8'>{lw or '(empty)'}</span> → <tt>Auto:</tt> <span style='background:#e8f5e9'>{rw or '(empty)'}</span>"
                    for lw, rw in diff_pairs[:20]
                ]
            sections.append((f"Page {i+1}", issues))
    rows = "".join(
        f"<div style='margin-bottom:14px;border:1px solid #ddd;border-radius:6px;overflow:hidden;'>"
        f"<div style='background:{'#fde8e8' if any('color:#c0392b' in x for x in iss) else '#e8f5e9'};padding:8px 12px;font-weight:bold;color:{'#c0392b' if any('color:#c0392b' in x for x in iss) else '#196F3D'};font-size:14px;'>{lbl}</div>"
        f"<div style='padding:8px 14px;font-family:monospace;font-size:12px;line-height:1.8;'>{''.join(f'<div>{x}</div>' for x in iss)
                                                                                               }</div></div>"
        for lbl, iss in sections
    )
    html = f"<html><body style='font-family:Segoe UI,Arial,sans-serif;background:#f8f9fa;color:#333;padding:16px;'>{rows}</body></html>"
    return {"html": html, "differences": [], "total_pages": total_pages}


# ══════════════════════════════════════════════════════════════════════════════
# PGTA ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/pgta/upload-cnv")
async def pgta_upload_cnv(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ('.png', '.jpg', '.jpeg'):
        return {"error": "Only PNG/JPG images allowed"}
    unique_name = str(uuid.uuid4()) + ext
    save_path = os.path.join(PGTA_CNV_DIR, unique_name)
    with open(save_path, "wb") as f:
        f.write(await file.read())
    return {"path": save_path, "name": unique_name, "url": f"/pgta/cnv-image/{unique_name}"}


@app.get("/pgta/cnv-image/{filename}")
def pgta_get_cnv_image(filename: str):
    path = os.path.join(PGTA_CNV_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(404, "Image not found")
    return FileResponse(path)


@app.post("/pgta/preview")
async def pgta_preview(request: Request):
    try:
        data = await request.json()
        file_id = str(uuid.uuid4()) + ".pdf"
        filepath = os.path.join(TEMP_DIR, file_id)
        embryos = data.get("embryos_data", [])
        embryos, tmp = _resolve_cnv_images(embryos)
        PGTAReportTemplate().generate_pdf(
            filepath, data.get("patient_data", {}), embryos,
            show_logo=data.get("show_logo", True), show_grid=data.get("show_grid", False)
        )
        for p in tmp:
            try:
                os.remove(p)
            except:
                pass
        return {"preview_url": f"/pgta/preview-file/{file_id}"}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": str(e)}


@app.get("/pgta/preview-file/{filename}")
def pgta_preview_file(filename: str):
    path = os.path.join(TEMP_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(404, "Preview not found")
    return FileResponse(path, media_type="application/pdf")


@app.post("/pgta/generate")
async def pgta_generate(request: Request, background_tasks: BackgroundTasks):
    try:
        data = await request.json()
        patient_info = data.get("patient_info") or data.get("patient_data", {})
        embryos = data.get("embryos") or data.get("embryos_data", [])
        options = data.get("options", {})
        show_logo = options.get("show_logo", True)
        show_grid = options.get("show_grid", False)
        formats = options.get("formats", ["pdf"])

        embryos, tmp_cnv = _resolve_cnv_images(embryos)

        def _parts(raw):
            return [p for p in re.sub(r'[^a-zA-Z0-9 ]', '', str(raw or '')).strip().split() if p]

        p_parts = _parts(patient_info.get("patient_name", ""))
        s_parts = _parts(patient_info.get("spouse_name",  ""))
        p_first = p_parts[0].upper() if p_parts else "UNKNOWN"
        p_init = p_parts[-1][0].upper() if len(p_parts) > 1 else (p_parts[0]
                                                                  [0].upper() if p_parts else "X")
        s_raw = str(patient_info.get("spouse_name", "") or "").strip().upper()
        if s_parts and s_raw not in ("WO", "W/O", "NA", "N/A", ""):
            s_first = s_parts[0].upper()
            s_init = s_parts[-1][0].upper() if len(
                s_parts) > 1 else s_parts[0][0].upper()
            name_seg = f"{p_first}_{s_first}_{p_init}_{s_init}"
        else:
            name_seg = f"{p_first}_{p_init}"

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        logo_tag = "withlogo" if show_logo else "withoutlogo"
        base_fn = f"PGTA_{name_seg}_{ts}_{logo_tag}"
        results = {}

        if "pdf" in formats:
            fn = base_fn + ".pdf"
            fp = os.path.join(PGTA_REPORT_DIR, fn)
            PGTAReportTemplate().generate_pdf(fp, patient_info, embryos,
                                              show_logo=show_logo, show_grid=show_grid)
            if upload_pgta_file:
                background_tasks.add_task(_upload_in_background, fp, fn)
            results["pdf"] = {"file": fn, "url": f"/reports-pgta/{fn}", "download_url": f"/pgta/download/{fn}"}

        if "docx" in formats:
            fn = base_fn + ".docx"
            fp = os.path.join(PGTA_REPORT_DIR, fn)
            PGTADocxGenerator(assets_dir=os.path.join(BASE_DIR, "assets", "pgta")).generate_docx(
                fp, patient_info, embryos, show_logo=show_logo, show_grid=show_grid)
            if upload_pgta_file:
                background_tasks.add_task(_upload_in_background, fp, fn)
            results["docx"] = {"file": fn, "url": f"/reports-pgta/{fn}", "download_url": f"/pgta/download/{fn}"}

        for p in tmp_cnv:
            try:
                os.remove(p)
            except:
                pass

        return {"status": "success", "results": results}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": str(e)}


@app.get("/pgta/download/{filename}")
def pgta_download(filename: str):
    """Force a real file download (Content-Disposition: attachment) instead of letting the
    browser open PDFs inline in its own viewer."""
    path = os.path.join(PGTA_REPORT_DIR, os.path.basename(filename))
    if not os.path.exists(path):
        raise HTTPException(404, "File not found")
    media_type = ("application/pdf" if filename.lower().endswith(".pdf")
                  else "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    return FileResponse(path, media_type=media_type, filename=filename)


# ── PGTA Excel parsing ─────────────────────────────────────────────────────────

async def _parse_pgta_excel_core(contents: bytes):
    xl = pd.ExcelFile(io.BytesIO(contents))
    sheets = xl.sheet_names
    sheets_lower = [s.lower().strip() for s in sheets]

    def clean_val(row, keys, default=''):
        if isinstance(keys, str):
            keys = [keys]
        for k in keys:
            if k in row:
                v = row[k]
                if pd.isna(v):
                    continue
                s = str(v).strip(' \t\r\f\v')
                if s.lower() in ('nan', 'none', 'nat', 'null'):
                    continue
                if s:
                    return s
        return default

    def fmt_date(val):
        if not val:
            return ""
        s = str(val).split(' ')[0]
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y", "%d.%m.%Y"):
            try:
                return datetime.strptime(s, fmt).strftime("%d-%m-%Y")
            except:
                continue
        return s.replace('/', '-')

    def norm(s):
        s = str(s or '').upper().strip()
        for pfx in ('MRS.', 'MR.', 'SMT.', 'DR.', 'MS.', 'MISS.', 'PROF.', 'R.', 'S.', 'K.', 'M.', 'D.', 'P.', 'A.', 'B.', 'C.', 'G.', 'H.', 'J.', 'L.', 'N.', 'T.', 'V.', 'W.'):
            if s.startswith(pfx):
                s = s[len(pfx):].strip()
                break
        s = re.sub(r'\([^)]*\)', '', s)
        return re.sub(r'[^A-Z0-9]', '', s)

    det_i = next((i for i, s in enumerate(
        sheets_lower) if 'detail' in s), None)
    sum_i = next((i for i, s in enumerate(
        sheets_lower) if 'summary' in s), None)
    if det_i is None and sum_i is None and len(sheets) >= 2:
        det_i, sum_i = 0, 1
    elif det_i is None and len(sheets) == 1:
        sum_i = 0

    det_df = sum_df = None
    if det_i is not None:
        det_df = xl.parse(sheets[det_i])
        det_df.columns = [str(c).strip() for c in det_df.columns]
    if sum_i is not None:
        try:
            df_full = xl.parse(sheets[sum_i], header=None)
            hdr = next((r for r, row in df_full.iterrows()
                        if any('sample name' in str(v).lower() for v in row.values)), 0)
            sum_df = xl.parse(sheets[sum_i], header=hdr)
            sum_df.columns = [str(c).strip() for c in sum_df.columns]
        except Exception:
            sum_df = xl.parse(sheets[sum_i])
            sum_df.columns = [str(c).strip() for c in sum_df.columns]

    patient_map, patients = {}, []

    if det_df is not None:
        for _, row in det_df.iterrows():
            name = clean_val(row, ['Patient Name', 'patient_name', 'Name'])
            if not name:
                continue
            pid = clean_val(
                row, ['Sample ID', 'Patient ID', 'sample_id', 'PIN', 'pin'])
            b_date = fmt_date(
                clean_val(row, ['Date of Biopsy', 'Biopsy Date']))
            r_date = fmt_date(clean_val(
                row, ['Date Sample Received', 'Receipt Date', 'Sample Receipt Date']))
            p = {
                "patient_name":        name,
                "sample_number":       clean_val(row, ['Sample Number', 'Sample No', 'Sample No.', 'Accession Number', 'Acc. No.'], ''),
                "hospital_clinic":     clean_val(row, ['Center name', 'Center Name', 'Hospital', 'Clinic']),
                "biopsy_date":         b_date,
                "sample_receipt_date": r_date,
                "biopsy_performed_by": clean_val(row, ['EMBRYOLOGIST NAME', 'Embryologist Name', 'Biologist']),
                "spouse_name":         clean_val(row, ['Spouse Name', 'Husband Name', 'Partner Name', 'spouse_name'], 'w/o'),
                "pin":                 pid,
                "age":                 clean_val(row, ['Age', 'age', 'Patient Age']),
                "referring_clinician": clean_val(row, ['Referring Clinician', 'referring_clinician', 'Doctor']),
                "specimen":            clean_val(row, ['Specimen', 'Specimen Type', 'Sample Type'], 'DAY 5 TROPHECTODERM BIOPSY'),
                "report_date":         datetime.now().strftime('%d-%m-%Y'),
                "indication":          clean_val(row, ['Indication', 'indication', 'Clinical Indication']),
                "embryos":             []
            }
            patient_map[norm(pid)] = p
            patient_map[norm(name)] = p
            patients.append(p)

    if sum_df is not None:
        for _, row in sum_df.iterrows():
            sname = clean_val(
                row, ['Sample name', 'Sample Name', 'sample_name', 'Sample ID'])
            if not sname:
                continue
            emb = {
                "embryo_id":           sname,
                "result_summary":      clean_val(row, ['Result', 'result', 'Summary']),
                "interpretation":      clean_val(row, ['Conclusion', 'Interpretation', 'interpretation']),
                "mtcopy":              clean_val(row, ['MTcopy', 'MT Copy', 'mtcopy', 'MT']),
                "autosomes":           clean_val(row, ['AUTOSOMES', 'Autosomes', 'autosomes', 'Aneuploidy']),
                "sex_chromosomes":     clean_val(row, ['SEX', 'Sex Chromosomes', 'sex_chromosomes', 'Sex'], 'Normal'),
                "result_description":  clean_val(row, ['Result', 'result_description']),
                "chromosome_statuses": {},
                "mosaic_percentages":  {},
                "inconclusive_comment": ""
            }
            ns = norm(sname)
            matched = next(
                (p for k, p in patient_map.items() if k and k in ns), None)
            if matched:
                matched["embryos"].append(emb)
            elif patients:
                patients[-1]["embryos"].append(emb)

    return {"patients": patients, "sheet_names": sheets}


@app.post("/pgta/parse-excel")
async def pgta_parse_excel(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        return await _parse_pgta_excel_core(contents)
    except Exception as e:
        return {"error": str(e)}


@app.post("/pgta/parse-excel-bulk")
async def pgta_parse_excel_bulk(files: List[UploadFile] = File(...)):
    try:
        excel_file = next((f for f in files if os.path.splitext(
            f.filename or "")[1].lower() in ('.xlsx', '.xls')), None)
        image_files = [f for f in files if os.path.splitext(
            f.filename or "")[1].lower() in ('.png', '.jpg', '.jpeg')]
        if not excel_file:
            return {"error": "No Excel file found"}
        data = await _parse_pgta_excel_core(await excel_file.read())
        patients = data.get("patients", [])
        all_embs = [e for p in patients for e in p.get("embryos", [])]
        mapped = auto_map_cnvs(all_embs, [f.filename for f in image_files])
        for f in image_files:
            if any(e.get("cnv_image_name") == f.filename for e in all_embs):
                content = await f.read()
                uname = str(uuid.uuid4()) + \
                    os.path.splitext(f.filename)[1].lower()
                spath = os.path.join(PGTA_CNV_DIR, uname)
                with open(spath, "wb") as out:
                    out.write(content)
                for e in all_embs:
                    if e.get("cnv_image_name") == f.filename:
                        e["cnv_image_path"] = spath
                        e["cnv_image_url"] = f"/pgta/cnv-image/{uname}"
                        e["cnv_image_b64"] = "data:image/png;base64," + \
                            base64.b64encode(content).decode()
        return {"patients": patients, "sheet_names": data.get("sheet_names", []), "mapped_count": mapped}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": str(e)}


# ── PGTA Drafts ────────────────────────────────────────────────────────────────

@app.post("/pgta/draft/save")
async def pgta_save_draft(request: Request):
    try:
        body = await request.json()
        patient = body.get("patient", {})
        pname = re.sub(r'[^a-zA-Z0-9 ]', '', str(patient.get("patient_name",
                       "draft"))).replace(" ", "_").strip() or "draft"
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        fname = f"pgta_bulk_draft_{pname}_{ts}.json"
        fpath = os.path.join(PGTA_DRAFT_DIR, fname)
        with open(fpath, "w", encoding="utf-8") as f:
            json.dump({"patients": [patient], "_type": "pgta_bulk_draft",
                      "_savedAt": datetime.now().isoformat()}, f, indent=2)
        return {"status": "saved", "filename": fname}
    except Exception as e:
        return {"status": "error", "error": str(e)}


@app.get("/pgta/draft/list")
def pgta_list_drafts():
    try:
        files = sorted([f for f in os.listdir(PGTA_DRAFT_DIR)
                       if f.endswith(".json")], reverse=True)
        return {"files": files}
    except Exception as e:
        return {"files": [], "error": str(e)}


@app.delete("/pgta/draft/delete/{filename}")
def pgta_delete_draft(filename: str):
    try:
        fp = os.path.join(PGTA_DRAFT_DIR, os.path.basename(filename))
        if os.path.exists(fp):
            os.remove(fp)
            return {"status": "deleted"}
        return {"status": "not_found"}
    except Exception as e:
        return {"status": "error", "error": str(e)}


# ── PGTA Compare ───────────────────────────────────────────────────────────────

@app.post("/pgta/compare")
async def pgta_compare(manual: UploadFile = File(...), automated: UploadFile = File(...)):
    def extract(data, fname):
        if fname.lower().endswith(".pdf"):
            try:
                lines = []
                with pdfplumber.open(io.BytesIO(data)) as pdf:
                    for pg in pdf.pages:
                        t = pg.extract_text()
                        if t:
                            lines.extend(t.splitlines())
                return [l.strip() for l in lines if l.strip()]
            except Exception as e:
                return [f"[PDF error: {e}]"]
        elif fname.lower().endswith(".docx"):
            try:
                from docx import Document
                doc = Document(io.BytesIO(data))
                return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            except Exception as e:
                return [f"[DOCX error: {e}]"]
        return ["[Unsupported format]"]

    mb, ab = await manual.read(), await automated.read()
    ml, al = extract(mb, manual.filename), extract(ab, automated.filename)
    d = difflib.HtmlDiff(wrapcolumn=80)
    html = d.make_table(ml, al, fromdesc=f"Manual — {manual.filename}",
                        todesc=f"Automated — {automated.filename}", context=True, numlines=3)
    changes = [(t, i1, i2, j1, j2) for t, i1, i2, j1, j2 in difflib.SequenceMatcher(
        None, ml, al).get_opcodes() if t != "equal"]
    return {"html_diff": html, "total_changes": len(changes), "match": len(changes) == 0}


# ── PGTA local storage list ────────────────────────────────────────────────────

@app.get("/pgta/storage/list")
async def pgta_storage_list(path: str = ""):
    try:
        files = sorted(
            [f for f in os.listdir(PGTA_REPORT_DIR) if not f.startswith(".")],
            reverse=True,
        )
        items = [{"name": f, "id": None, "metadata": None} for f in files]
        return {"items": items, "path": path}
    except Exception as exc:
        raise HTTPException(500, str(exc))


# ── Folder-picker helper (Windows server only) ────────────────────────────────

@app.get("/open-folder-dialog")
async def open_folder_dialog():
    import asyncio
    try:
        import subprocess
        ps = r"C:\Windows\System32\WindowsPowerShell\v1.0\powershell.exe"
        script = (
            "Add-Type -AssemblyName System.Windows.Forms;"
            "Add-Type -AssemblyName System.Drawing;"
            "[System.Windows.Forms.Application]::EnableVisualStyles();"
            "$owner = New-Object System.Windows.Forms.Form;"
            "$owner.TopMost = $true;"
            "$owner.Size = New-Object System.Drawing.Size(1,1);"
            "$owner.StartPosition = 'CenterScreen';"
            "$owner.Show();"
            "$owner.Activate();"
            "$d = New-Object System.Windows.Forms.FolderBrowserDialog;"
            "$d.Description = 'Select Output Folder';"
            "$d.ShowNewFolderButton = $true;"
            "$d.AutoUpgradeEnabled = $true;"
            "if ($d.ShowDialog($owner) -eq 'OK') { Write-Output $d.SelectedPath };"
            "$owner.Dispose();"
        )
        res = await asyncio.get_event_loop().run_in_executor(
            None, lambda: subprocess.run([ps, "-NoProfile", "-Command", script],
                                         capture_output=True, text=True, timeout=120))
        return {"path": res.stdout.strip()}
    except Exception as e:
        return {"error": str(e)}


# ── Health check ───────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {"status": "ok", "mysql": _mysql_enabled, "timestamp": datetime.now().isoformat()}


# ══════════════════════════════════════════════════════════════════════════════
# NIPT ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

NIPT_REPORT_DIR = os.path.join(BASE_DIR, "reports-nipt")
NIPT_DRAFT_DIR = os.path.join(BASE_DIR, "drafts", "NIPT")
os.makedirs(NIPT_REPORT_DIR, exist_ok=True)
os.makedirs(NIPT_DRAFT_DIR,  exist_ok=True)

# Mount NIPT reports for static access
app.mount("/reports-nipt", StaticFiles(directory=NIPT_REPORT_DIR),
          name="reports-nipt")


@app.get("/nipt/download/{filename}")
def nipt_download_file(filename: str):
    """Force-download a NIPT report with Content-Disposition: attachment."""
    safe = os.path.basename(filename)
    fp = os.path.join(NIPT_REPORT_DIR, safe)
    if not os.path.exists(fp):
        raise HTTPException(404, "File not found")
    media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" \
        if safe.endswith(".docx") else "application/pdf"
    return FileResponse(fp, media_type=media,
                        headers={"Content-Disposition": f'attachment; filename="{safe}"'})


# ── Normalisation helpers (mirrors desktop nipt_report_generator.py) ──────────

_HOSPITAL_ACRONYMS = {
    'AIIMS', 'JIPMER', 'NIMHANS', 'PGIMER', 'SGPGI', 'NIMS', 'AFMC', 'CMC', 'MGM', 'KMC',
    'GMC', 'RMC', 'SMS', 'IMS', 'PGI', 'KIMS', 'RIMS', 'MIMS', 'SIMS', 'JSS', 'SDM', 'ESI',
    'ESIC', 'AMC', 'SRMC', 'TNMC', 'IVF', 'IUI', 'ICSI', 'FET', 'ART', 'ENT', 'ICU', 'NICU',
    'PICU', 'OPD', 'IPD', 'MRI', 'CT',
}


def _title_case_words(value: str) -> str:
    text = " ".join(w[:1].upper() + w[1:].lower()
                    for w in str(value or "").strip().split())
    return re.sub(r"\b(Mr|Mrs|Ms|Dr)\.\s*([a-z])",
                  lambda m: f"{m.group(1)}. {m.group(2).upper()}", text)


def _fmt_hospital(value: str) -> str:
    def _w(w):
        u = w.upper()
        if u in _HOSPITAL_ACRONYMS:
            return u
        if w.isalpha() and not any(c in 'aeiouAEIOU' for c in w):
            return u
        return w.capitalize()
    return " ".join(_w(t) for t in str(value or "").strip().split())


def _norm_patient(p: dict) -> dict:
    for k in ("name", "clinician"):
        if k in p:
            p[k] = _title_case_words(p.get(k, ""))
    if "hospital" in p:
        p["hospital"] = _fmt_hospital(p.get("hospital", ""))
    if "clinician_qual" in p:
        p["clinician_qual"] = re.sub(r'[A-Za-z]+', lambda m: m.group().upper(),
                                     str(p.get("clinician_qual", "")))
    return p


def _nipt_base_filename(name: str, with_logo: bool) -> str:
    n = _title_case_words(name) or "Patient"
    n = re.sub(r'[<>:"/\\|?*]+', "", n)
    n = re.sub(r"\s+", "_", n).strip("_.") or "Patient"
    return f"{n}_NIPT_Report_{'with_logo' if with_logo else 'without_logo'}"


def _safe_float(v):
    if v is None:
        return 0.0
    try:
        f = float(v)
        return 0.0 if math.isnan(f) else f
    except:
        return 0.0


def _norm_id(v) -> str:
    s = str(v).strip()
    try:
        f = float(s)
        if f == int(f):
            return str(int(f))
    except:
        pass
    return s


# ── Preview ───────────────────────────────────────────────────────────────────

@app.post("/nipt/preview")
async def nipt_preview(request: Request):
    if not _nipt_ok:
        return {"error": "NIPT generators not loaded on server"}
    try:
        data = await request.json()
        file_id = str(uuid.uuid4()) + ".pdf"
        filepath = os.path.join(TEMP_DIR, file_id)
        p_info = _norm_patient(dict(data.get("patient_data", {})))
        z_scores = {k: _safe_float(v)
                    for k, v in data.get("z_scores", {}).items()}
        show_logo = bool(data.get("show_logo", True))
        NIPTReportTemplate(filepath).generate(
            p_info, z_scores, with_logo=show_logo)
        return {"preview_url": f"/nipt/preview-file/{file_id}"}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": str(e)}


@app.get("/nipt/preview-file/{filename}")
def nipt_preview_file(filename: str):
    path = os.path.join(TEMP_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(404, "Preview not found")
    return FileResponse(path, media_type="application/pdf")


# ── Generate (single patient) ─────────────────────────────────────────────────

@app.post("/nipt/generate")
async def nipt_generate(request: Request):
    if not _nipt_ok:
        return {"error": "NIPT generators not loaded on server"}
    try:
        data = await request.json()
        p_info = _norm_patient(dict(data.get("patient_data", {})))
        z_scores = {k: _safe_float(v)
                    for k, v in data.get("z_scores", {}).items()}
        options = data.get("options", {})
        show_logo = bool(options.get("show_logo", True))
        formats = options.get("formats", ["pdf"])

        base = _nipt_base_filename(p_info.get("name", ""), show_logo)
        results = {}

        if "pdf" in formats:
            fn = base + ".pdf"
            fp = os.path.join(NIPT_REPORT_DIR, fn)
            NIPTReportTemplate(fp).generate(
                p_info, z_scores, with_logo=show_logo)
            results["pdf"] = {"file": fn, "url": f"/reports-nipt/{fn}"}

        if "docx" in formats:
            fn = base + ".docx"
            fp = os.path.join(NIPT_REPORT_DIR, fn)
            NIPTDocxGenerator(fp).generate(
                p_info, z_scores, with_logo=show_logo)
            results["docx"] = {"file": fn, "url": f"/reports-nipt/{fn}"}

        return {"status": "success", "results": results}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": str(e)}


# ── Generate batch (all patients in one request) ───────────────────────────────

@app.post("/nipt/generate-batch")
async def nipt_generate_batch(request: Request):
    if not _nipt_ok:
        return {"error": "NIPT generators not loaded on server"}
    try:
        data     = await request.json()
        patients = data.get("patients", [])
        options  = data.get("options", {})
        show_logo = bool(options.get("show_logo", True))
        formats   = options.get("formats", ["pdf"])

        def _gen_one(pat):
            p_info   = _norm_patient(dict(pat))
            z_scores = {k: _safe_float(v) for k, v in pat.items()
                        if k.startswith("chr")}
            base = _nipt_base_filename(p_info.get("name", ""), show_logo)
            out  = {}
            if "pdf" in formats:
                fn = base + ".pdf"
                fp = os.path.join(NIPT_REPORT_DIR, fn)
                NIPTReportTemplate(fp).generate(p_info, z_scores, with_logo=show_logo)
                out["pdf"] = {"file": fn, "url": f"/reports-nipt/{fn}"}
            if "docx" in formats:
                fn = base + ".docx"
                fp = os.path.join(NIPT_REPORT_DIR, fn)
                NIPTDocxGenerator(fp).generate(p_info, z_scores, with_logo=show_logo)
                out["docx"] = {"file": fn, "url": f"/reports-nipt/{fn}"}
            return {"name": p_info.get("name", ""), "results": out}

        import concurrent.futures
        loop    = asyncio.get_event_loop()
        results = []
        errors  = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=4) as pool:
            futs = {loop.run_in_executor(pool, _gen_one, pat): pat for pat in patients}
            for fut in asyncio.as_completed(futs):
                try:
                    results.append(await fut)
                except Exception as e:
                    errors.append(str(e))

        return {"status": "success", "results": results, "errors": errors}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": str(e)}


# ── Parse Excel (batch) ────────────────────────────────────────────────────────

@app.post("/nipt/parse-excel")
async def nipt_parse_excel(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        xls = pd.ExcelFile(io.BytesIO(contents))
        sname = "Sheet1" if "Sheet1" in xls.sheet_names else xls.sheet_names[0]
        df = pd.read_excel(xls, sname)
        df.columns = [str(c).strip() for c in df.columns]

        # QC filter
        if "QC" in df.columns:
            df = df[df["QC"].astype(str).str.strip(
            ).str.lower() == "pass"].reset_index(drop=True)

        # Sheet2 Z-scores merge
        if "Sheet2" in xls.sheet_names:
            df_z = pd.read_excel(xls, "Sheet2")
            df_z.columns = [str(c).strip() for c in df_z.columns]
            col_rename = {}
            for col in df_z.columns:
                norm = str(col).strip().lower().replace(
                    ' ', '').replace('_', '').replace('-', '')
                for i in range(1, 23):
                    if norm in (f"chr{i}", f"chromosome{i}", f"chrom{i}"):
                        col_rename[col] = f"chr{i}"
                        break
                else:
                    if norm in ('chrx', 'chromosomex', 'chromx'):
                        col_rename[col] = 'chrX'
                    elif norm in ('fetaldna', 'fetalfraction', 'fetaldnafraction',
                                  'fetaldna%', 'fetalfraction%', 'ff'):
                        col_rename[col] = 'Fetal DNA'
            if col_rename:
                df_z = df_z.rename(columns=col_rename)

            id1 = next(
                (c for c in ["Sample ID", "Sample Name"] if c in df.columns),   None)
            id2 = next(
                (c for c in ["Sample ID", "Sample"] if c in df_z.columns), None)
            if id1 and id2:
                df[id1] = df[id1].apply(_norm_id)
                df_z[id2] = df_z[id2].apply(_norm_id)
                matched_ids = set(df[id1]) & set(df_z[id2])
                unmatched_pos = [i for i, v in enumerate(
                    df[id1]) if v not in matched_ids]
                leftover_z = df_z[~df_z[id2].isin(
                    matched_ids)].reset_index(drop=True)
                df = pd.merge(df, df_z, left_on=id1, right_on=id2,
                              how="left").reset_index(drop=True)
                if len(unmatched_pos) > 0 and len(unmatched_pos) == len(leftover_z):
                    z_cols = [c for c in df_z.columns if c != id2]
                    for pos, s1p in enumerate(unmatched_pos):
                        for col in z_cols:
                            if col in df.columns:
                                df.at[s1p, col] = leftover_z.at[pos, col]

        raw = df.to_dict("records")
        patients = []
        for row in raw:
            ff = _safe_float(row.get("Fetal DNA", row.get("FF", 0)))
            ff_pct = ff * 100 if ff < 1 else ff

            def _s(v):
                if v is None:
                    return ""
                try:
                    if pd.isna(v):
                        return ""
                except:
                    pass
                s = str(v).replace(" 00:00:00", "").split(" /")[0].strip()
                return "" if s.lower() in ("nan", "nat", "none", "null") else s

            p = {
                "name":            _title_case_words(_s(row.get("Patient Name", row.get("Sample Name", "")))),
                "pin":             _s(row.get("PIN", row.get("Sample Name", ""))),
                "dob":             _s(row.get("Date of Birth", row.get("DOB", row.get("Date of birth/Age", "")))),
                "age":             _s(row.get("Age", "")),
                "ga":              _s(row.get("Gestational Age", "")),
                "sample_id":       _s(row.get("Sample ID", row.get("Sample Name", ""))),
                "collection_date": _s(row.get("Collection date", row.get("Col Date", ""))),
                "received_date":   _s(row.get("Received date", row.get("Rec Date", ""))),
                "preg_status":     _s(row.get("Pregnancy status", row.get("Status", ""))),
                "preg_type":       _s(row.get("Pregnancy type", "")),
                "clinician":       _title_case_words(_s(row.get("Referring Clinician", row.get("Ref Doctor", "")))),
                "clinician_qual":  _s(row.get("Qualification", row.get("Clinician Qualification", ""))),
                "hospital":        _fmt_hospital(_s(row.get("Hospital", ""))),
                "indication":      _s(row.get("Indication", "")),
                "specimen":        _s(row.get("Specimen", "")),
                "ff":              str(ff_pct),
            }
            for i in range(1, 23):
                val = _safe_float(row.get(f"chr{i}", 0))
                p[f"chr{i}"] = f"{math.trunc(val * 100) / 100:.2f}"
            p["chrX"] = f"{math.trunc(_safe_float(row.get('chrX', 0)) * 100) / 100:.2f}"
            patients.append(p)

        return {"patients": patients, "count": len(patients)}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": str(e), "patients": []}


# ── Drafts ─────────────────────────────────────────────────────────────────────

@app.post("/nipt/draft/save")
async def nipt_save_draft(request: Request):
    try:
        body = await request.json()
        pname = re.sub(r'[^a-zA-Z0-9 ]', '',
                       str(body.get("patient_details", {}).get("name", "draft"))
                       ).replace(" ", "_").strip() or "draft"
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        fname = f"nipt_draft_{pname}_{ts}.json"
        with open(os.path.join(NIPT_DRAFT_DIR, fname), "w", encoding="utf-8") as f:
            json.dump(body, f, indent=2)
        return {"status": "saved", "filename": fname}
    except Exception as e:
        return {"status": "error", "error": str(e)}


@app.get("/nipt/draft/list")
def nipt_list_drafts():
    try:
        files = sorted([f for f in os.listdir(NIPT_DRAFT_DIR)
                       if f.endswith(".json")], reverse=True)
        return {"files": files}
    except Exception as e:
        return {"files": [], "error": str(e)}


@app.get("/nipt/draft/load/{filename}")
def nipt_load_draft(filename: str):
    fpath = os.path.join(NIPT_DRAFT_DIR, os.path.basename(filename))
    if not os.path.exists(fpath):
        return {"data": None}
    with open(fpath, encoding="utf-8") as f:
        return {"data": json.load(f)}


@app.delete("/nipt/draft/delete/{filename}")
def nipt_delete_draft(filename: str):
    fp = os.path.join(NIPT_DRAFT_DIR, os.path.basename(filename))
    if os.path.exists(fp):
        os.remove(fp)
        return {"status": "deleted"}
    return {"status": "not_found"}


# ── Compare PDFs (file mode — reuse existing pdfplumber logic) ─────────────────

@app.post("/nipt/compare-pdf")
async def nipt_compare_pdf(manual: UploadFile = File(...), automated: UploadFile = File(...)):
    def extract(data, fname):
        if fname.lower().endswith(".pdf"):
            try:
                pages, lines = [], []
                with pdfplumber.open(io.BytesIO(data)) as pdf:
                    for pg in pdf.pages:
                        t = pg.extract_text() or ""
                        pages.append(_norm(t))
                        lines.extend([l.strip()
                                     for l in t.splitlines() if l.strip()])
                return pages, lines
            except Exception as e:
                return [], [f"[PDF error: {e}]"]
        elif fname.lower().endswith(".docx"):
            try:
                from docx import Document
                doc = Document(io.BytesIO(data))
                lines = [p.text.strip()
                         for p in doc.paragraphs if p.text.strip()]
                return [], lines
            except Exception as e:
                return [], [f"[DOCX error: {e}]"]
        return [], ["[Unsupported format]"]

    mb, ab = await manual.read(), await automated.read()
    mp, ml = extract(mb, manual.filename)
    ap, al = extract(ab, automated.filename)
    d = difflib.HtmlDiff(wrapcolumn=80)
    html = d.make_table(ml, al,
                        fromdesc=f"Manual — {manual.filename}",
                        todesc=f"Automated — {automated.filename}",
                        context=True, numlines=3)
    changes = [(t, i1, i2, j1, j2) for t, i1, i2, j1, j2 in
               difflib.SequenceMatcher(None, ml, al).get_opcodes() if t != "equal"]
    total_pages = max(len(mp), len(ap)) if mp or ap else 1
    return {
        "html_diff": html,
        "total_changes": len(changes),
        "match": len(changes) == 0,
        "total_pages": total_pages,
        "manual_file": manual.filename,
        "auto_file": automated.filename
    }


# ── Compare directories (dir mode) ────────────────────────────────────────────

@app.post("/nipt/compare-dir")
async def nipt_compare_dir(request: Request):
    try:
        body = await request.json()
        m_dir = body.get("manual_dir", "")
        a_dir = body.get("auto_dir", "")
        if not os.path.isdir(m_dir) or not os.path.isdir(a_dir):
            return {"error": "One or both directories do not exist on the server"}
        mf = {f for f in os.listdir(m_dir) if f.lower().endswith(".pdf")}
        af = {f for f in os.listdir(a_dir) if f.lower().endswith(".pdf")}
        only_m = sorted(mf - af)
        only_a = sorted(af - mf)
        matched = sorted(mf & af)
        rows = [
            f"<h2 style='color:#1F497D'>Directory Comparison Report</h2>"
            f"<p><b>Manual:</b> {m_dir}<br><b>Automated:</b> {a_dir}</p>"
            f"<p>Total: {len(mf | af)} files | Matched: {len(matched)} | "
            f"Manual only: {len(only_m)} | Auto only: {len(only_a)}</p><hr>"
        ]
        if only_m:
            rows.append("<h3 style='color:#d97706'>Only in Manual:</h3><ul>" +
                        "".join(f"<li>{f}</li>" for f in only_m) + "</ul>")
        if only_a:
            rows.append("<h3 style='color:#7c3aed'>Only in Automated:</h3><ul>" +
                        "".join(f"<li>{f}</li>" for f in only_a) + "</ul>")
        if matched:
            rows.append(f"<h3 style='color:#16a34a'>Matched Pairs ({len(matched)}):</h3><ul>" +
                        "".join(f"<li><b>{f}</b> — matched by name</li>" for f in matched) + "</ul>")
        return {"html": "".join(rows)}
    except Exception as e:
        return {"error": str(e)}


# ── Storage list ───────────────────────────────────────────────────────────────

@app.get("/nipt/storage/list")
async def nipt_storage_list():
    try:
        files = sorted([f for f in os.listdir(NIPT_REPORT_DIR)
                       if not f.startswith(".")], reverse=True)
        return {"items": [{"name": f, "url": f"/reports-nipt/{f}"} for f in files]}
    except Exception as exc:
        raise HTTPException(500, str(exc))


# ── Karyotype report routes (web port of desktop Karyotype Report Generator) ──
try:
    from karyotype_api import register_karyotype_routes
    register_karyotype_routes(app)
    print("[KARYOTYPE] routes loaded")
except Exception as _karyo_err:
    print(f"[KARYOTYPE] routes not loaded: {_karyo_err}")

# ── HLA Typing report routes (web port of desktop HLA Report Generator) ────────
if _hla_ok:
    app.include_router(hla_router)
    _hla_fonts_dir = os.path.join(BASE_DIR, "assets", "hla", "fonts")
    if os.path.isdir(_hla_fonts_dir):
        app.mount("/hla-fonts", StaticFiles(directory=_hla_fonts_dir), name="hla-fonts")
    print("[HLA] routes loaded")


@app.get("/hla")
@app.get("/hla.html")
def hla_page():
    p = os.path.join(FRONTEND_DIR, "hla.html")
    if os.path.exists(p):
        return FileResponse(p, media_type="text/html",
                            headers={"Cache-Control": "no-cache, no-store, must-revalidate"})
    raise HTTPException(404, "hla.html not found")


@app.get("/hla.js")
def hla_js():
    p = os.path.join(FRONTEND_DIR, "hla.js")
    if os.path.exists(p):
        return FileResponse(p, media_type="application/javascript",
                            headers={"Cache-Control": "no-cache, no-store, must-revalidate"})
    raise HTTPException(404, "hla.js not found")
